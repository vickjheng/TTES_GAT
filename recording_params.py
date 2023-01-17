#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Tue Jan 10 11:17:43 2023

@author: dennischang
"""

from param import args
import datetime
import docx
import numpy as np
import matplotlib.pyplot as plt
from docx.shared import Cm  #加入可調整的 word 單位

episode = 3000
window_size = episode / 40


def move_avg(data):
    window = np.ones(int(window_size)) / float(window_size)

    return np.convolve(data, window, 'same')

def draw_env_record():
    env_record = {'success_rate': [], 'usage': [], 'reward': []}
    for key in env_record.keys():
        env_record[key].extend(move_avg(np.load(f'record/env/{key}_{episode}.npy', allow_pickle=True)))
    x = np.arange(len(env_record['success_rate']))
    fig, plot = plt.subplots(len(env_record))
    for idx, key in zip(range(len(env_record)), env_record.keys()):
        plot[idx].plot(x, env_record[key], 'r')
        plot[idx].set_xlim(int(window_size / 2), int(episode - (window_size / 2)))
        plot[idx].grid(True)
        plot[idx].set_title(f'{key}')
    fig.tight_layout()
    plt.savefig('env.png')
    y = move_avg(np.load(f'record/loss_{episode}.npy'))
    x = np.arange(len(y))
    plt.figure()
    plt.plot(x, y,color='r',label = 'loss')

    plt.xlim(int(window_size / 2), int(episode - (window_size / 2)))
    plt.title('Loss')
    #plt.axvline(args.exploration_end_episode, linestyle= '--',label='end_exploration')
    plt.grid(True)
    plt.legend(loc='best') 
    plt.savefig('loss.png')
    # plt.show()
   
draw_env_record()
doc = docx.Document()
version = datetime.datetime.now().strftime('[%m-%d %H:%M]')
title='-'*20+str(version)+'-'*20
doc.add_heading(str(title),level=2)

params = [str(args.episodes),str(args.capacity),str(args.epsilon_decay),str(args.batch_size),str(args.lr)]
labels = ['episodes','memory size','epsilon_decay','batch size','LR']
for param,label in zip(params,labels):
    doc.add_paragraph(f'{label}: {param}')
# doc.add_paragraph('(With LR decay)')
doc.add_picture('env.png', width=Cm(10))
doc.add_picture('loss.png', width=Cm(10))
doc.save('./resultrecording/'+str(version)+'_record.docx')